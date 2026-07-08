from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
FREEZE = ROOT / "outputs" / "pasd_phase4_paper_freeze"
INTEGRATED = FREEZE / "manuscript_integrated"
TABLE_DIR = FREEZE / "paper_tables"
FIGURE_DIR = FREEZE / "paper_figures"
SCI_FIGURE_DIR = FREEZE / "paper_figures_sci"
OUTPUT = INTEGRATED / "PASD_Core_checked_manuscript.docx"


CITATION_REPLACEMENTS = {
    "[REF: learned FWI baseline]": {
        "label": "learned FWI baseline",
        "replacement": "(Wu and Lin, 2020; Zhang et al., 2019)",
    },
    "[REF: OpenFWI or cross-family benchmark]": {
        "label": "OpenFWI or cross-family benchmark",
        "replacement": "(Deng et al., 2022)",
    },
    "[REF: neural FWI architectures]": {
        "label": "neural FWI architectures",
        "replacement": "(Wu and Lin, 2020; Zhang et al., 2019)",
    },
    "[REF: structural image metrics or edge-aware reconstruction]": {
        "label": "structural image metrics or edge-aware reconstruction",
        "replacement": "(Wang et al., 2004)",
    },
}


REFERENCES = [
    {
        "label": "Deng et al., 2022",
        "text": (
            "Deng, C., Feng, S., Wang, H., Zhang, X. et al. "
            "OpenFWI: Large-Scale Multi-Structural Benchmark Datasets for Full Waveform Inversion. "
            "Advances in Neural Information Processing Systems 35 (2022). "
            "https://doi.org/10.52202/068431-0435"
        ),
    },
    {
        "label": "Wu and Lin, 2020",
        "text": (
            "Wu, Y. and Lin, Y. InversionNet: An Efficient and Accurate Data-Driven Full Waveform Inversion. "
            "IEEE Transactions on Computational Imaging 6, 419-433 (2020). "
            "https://doi.org/10.1109/TCI.2019.2956866"
        ),
    },
    {
        "label": "Zhang et al., 2019",
        "text": (
            "Zhang, Z., Wu, Y., Zhou, Z. and Lin, Y. "
            "VelocityGAN: Subsurface Velocity Image Estimation Using Conditional Adversarial Networks. "
            "2019 IEEE Winter Conference on Applications of Computer Vision (WACV) (2019). "
            "https://doi.org/10.1109/WACV.2019.00080"
        ),
    },
    {
        "label": "Wang et al., 2004",
        "text": (
            "Wang, Z., Bovik, A. C., Sheikh, H. R. and Simoncelli, E. P. "
            "Image quality assessment: from error visibility to structural similarity. "
            "IEEE Transactions on Image Processing 13(4), 600-612 (2004). "
            "https://doi.org/10.1109/TIP.2003.819861"
        ),
    },
]


FIGURE_CAPTIONS = {
    "Figure_1_method_overview.png": "Figure 1. PASD-Core method overview.",
    "Figure_2_hybrid_bridge_attributes.png": "Figure 2. Physics-aligned hybrid bridge attributes.",
    "Figure_3_curvevel_median_comparison.png": "Figure 3. CurveVel-A median co-sample comparison.",
    "Figure_4_curvevel_hard_comparison.png": "Figure 4. CurveVel-A hard co-sample comparison.",
    "Figure_5_flatfault_median_comparison.png": "Figure 5. FlatFault-A median co-sample comparison.",
    "Figure_6_flatfault_hard_comparison.png": "Figure 6. FlatFault-A hard co-sample comparison.",
    "Figure_7_corrected_gradient_edge_comparison.png": "Figure 7. Corrected edge-gradient comparison.",
    "Figure_8_velocity_profiles.png": "Figure 8. Velocity-profile comparison.",
    "Figure_9_corrected_metric_distributions.png": "Figure 9. Corrected metric distributions.",
    "Figure_10_seed_bootstrap_summary.png": "Figure 10. Seed-level paired bootstrap summary.",
    "Figure_A1_geometry_attention_ablation.png": "Figure A1. Geometry-attention auxiliary ablation.",
}


SCI_FIGURE_CAPTIONS = {
    "Figure_SCI_1_evidence_summary.png": "SCI Figure 1. Cross-family evidence summary with directional improvement and seed-level bootstrap uncertainty.",
    "Figure_SCI_2_velocity_error_plate.png": "SCI Figure 2. Co-sample velocity and error plate with shared velocity, error, and error-reduction color scales.",
    "Figure_SCI_3_provenance_distributions.png": "SCI Figure 3. Target-isolation provenance and full-sample metric distributions.",
}


def set_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11.5)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text


def replace_citations(text: str) -> str:
    for placeholder, entry in CITATION_REPLACEMENTS.items():
        text = text.replace(placeholder, entry["replacement"])
    return text


def add_table_from_rows(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = clean_inline(rows[0][idx])
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(cells)]):
            cells[idx].text = clean_inline(value)


def add_markdown(doc: Document, markdown: str) -> None:
    lines = replace_citations(markdown).splitlines()
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            useful_rows = [row for row in table_rows if not all(set(col.strip()) <= {"-", ":"} for col in row)]
            add_table_from_rows(doc, useful_rows)
            table_rows = []

    for raw_line in lines:
        line = raw_line.rstrip()
        if line.startswith("|") and line.endswith("|"):
            parts = [part.strip() for part in line.strip("|").split("|")]
            table_rows.append(parts)
            continue
        flush_table()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(clean_inline(line[2:]), level=1)
        elif line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=2)
        elif line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(clean_inline(line[2:]), style="List Bullet")
        elif line.startswith("> "):
            paragraph = doc.add_paragraph(clean_inline(line[2:]))
            paragraph.style = doc.styles["Intense Quote"]
        else:
            doc.add_paragraph(clean_inline(line))
    flush_table()


def add_csv_table(doc: Document, csv_path: Path) -> None:
    with csv_path.open("r", encoding="utf-8", newline="") as handle:
        rows = list(csv.reader(handle))
    add_table_from_rows(doc, rows)


def add_title_page(doc: Document) -> None:
    doc.add_heading(
        "PASD-Core: Checked Manuscript and Academic-Search Verification",
        level=0,
    )
    doc.add_paragraph("Generated from the frozen Phase-4 evidence package.")
    doc.add_paragraph("Evidence source: results_summary.csv, Table 1-6, Figure 1-10/A1, and the integrated manuscript draft.")
    doc.add_paragraph("Scientific boundary: no model, metric, split, figure value, or result value was changed during this Word assembly.")
    doc.add_paragraph("Citation verification date: 2026-07-03.")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_verification_summary(doc: Document) -> None:
    doc.add_heading("Academic-Search Verification Summary", level=1)
    doc.add_paragraph(
        "The initial draft used explicit citation placeholders rather than invented references. "
        "The placeholders were checked against bibliographic metadata for representative learned FWI, "
        "OpenFWI benchmark, and structural similarity references. The replacements below are conservative "
        "and do not broaden the frozen scientific claim."
    )
    rows = [["Draft placeholder", "Replacement", "Verification status"]]
    for entry in CITATION_REPLACEMENTS.values():
        rows.append([entry["label"], entry["replacement"], "Metadata verified; suitable for bounded prior-work support."])
    add_table_from_rows(doc, rows)

    doc.add_heading("Review-Readiness Decision", level=2)
    doc.add_paragraph(
        "Current status: scientifically bounded and traceable, with formal references now inserted. "
        "Remaining submission work is target-journal formatting and any journal-specific bibliography style conversion."
    )
    doc.add_paragraph(
        "Do not add new experiments or modify frozen numerical evidence in this manuscript version.",
        style="List Bullet",
    )


def add_frozen_tables(doc: Document) -> None:
    doc.add_heading("Frozen Tables", level=1)
    for csv_path in sorted(TABLE_DIR.glob("Table_*.csv")):
        doc.add_heading(csv_path.stem.replace("_", " "), level=2)
        add_csv_table(doc, csv_path)


def add_figures(doc: Document) -> None:
    doc.add_heading("Frozen Figures", level=1)
    for name, caption in FIGURE_CAPTIONS.items():
        figure = FIGURE_DIR / name
        if not figure.exists():
            continue
        doc.add_heading(caption.split(".")[0], level=2)
        doc.add_picture(str(figure), width=Inches(6.2))
        doc.add_paragraph(caption)


def add_sci_figures(doc: Document) -> None:
    if not SCI_FIGURE_DIR.exists():
        return
    doc.add_heading("SCI-Grade Figure Upgrade Package", level=1)
    doc.add_paragraph(
        "These figures are generated from the same frozen Phase-4/Phase-3R evidence sources. "
        "They are intended as replacement candidates for the visually repetitive bar-chart figures; "
        "they do not introduce new metrics, new targets, or new model variants."
    )
    for name, caption in SCI_FIGURE_CAPTIONS.items():
        figure = SCI_FIGURE_DIR / name
        if not figure.exists():
            continue
        doc.add_heading(caption.split(".")[0], level=2)
        doc.add_picture(str(figure), width=Inches(6.2))
        doc.add_paragraph(caption)


def add_references(doc: Document) -> None:
    doc.add_heading("References", level=1)
    for ref in REFERENCES:
        doc.add_paragraph(ref["text"], style="List Number")


def main() -> None:
    manuscript = (INTEGRATED / "PASD_Core_manuscript_full_draft.md").read_text(encoding="utf-8")
    manuscript_body = manuscript.split("\n## References\n", 1)[0]
    audit = (INTEGRATED / "REVIEW_READINESS_AUDIT.md").read_text(encoding="utf-8")

    doc = Document()
    set_styles(doc)
    add_title_page(doc)
    add_verification_summary(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("Review Readiness Audit", level=1)
    add_markdown(doc, audit)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    doc.add_heading("Checked Manuscript Draft", level=1)
    add_markdown(doc, manuscript_body)
    add_references(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_frozen_tables(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_sci_figures(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_figures(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
